# -*- coding: utf-8 -*-
"""Convert early-course markdown briefs to DOCX + PDF."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

BASE = Path(__file__).resolve().parent
FILES = [
    BASE / "01-voice-actor-intro.md",
    BASE / "02-slovik-image-video-prompts.md",
    BASE / "03-voice-actor-trial-sounds.md",
]


def add_runs_with_code(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue
        segs = re.split(r"(\*\*[^*]+\*\*)", part)
        for seg in segs:
            if not seg:
                continue
            if seg.startswith("**") and seg.endswith("**"):
                run = paragraph.add_run(seg[2:-2])
                run.bold = True
            else:
                paragraph.add_run(seg)


def md_to_docx(md_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        code_lines = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith("|") and "|" in line[1:]:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in row):
                    i += 1
                    continue
                rows.append(row)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        txt = row[c_idx] if c_idx < len(row) else ""
                        add_runs_with_code(p, txt)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                doc.add_paragraph()
            continue
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_code(p, line.strip()[2:])
        elif re.match(r"^\d+\. ", line.strip()):
            p = doc.add_paragraph(style="List Number")
            add_runs_with_code(p, re.sub(r"^\d+\. ", "", line.strip()))
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            add_runs_with_code(p, line)
        i += 1

    flush_code()
    out = md_path.with_suffix(".docx")
    doc.save(out)
    print(f"DOCX {out}")
    return out


def docx_to_pdf_via_word(docx_paths: list[Path]) -> None:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        # fallback: Word.Application via dynamic Dispatch without pywin32 package name
        try:
            import comtypes.client  # type: ignore
        except ImportError:
            comtypes = None
        else:
            comtypes = True
        win32com = None
    else:
        comtypes = None

    # Prefer late-bound COM through win32com if available; else PowerShell-friendly ctypes is harder.
    # Use EnsureDispatch-like approach via Dispatch.
    try:
        import win32com.client as win32  # type: ignore
    except ImportError:
        print("NO_PYWIN32", file=sys.stderr)
        raise SystemExit(2)

    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        for docx_path in docx_paths:
            pdf_path = docx_path.with_suffix(".pdf")
            doc = word.Documents.Open(str(docx_path))
            # 17 = wdFormatPDF
            doc.SaveAs(str(pdf_path), FileFormat=17)
            doc.Close(False)
            print(f"PDF  {pdf_path}")
    finally:
        word.Quit()


def main() -> None:
    outs = [md_to_docx(p) for p in FILES]
    docx_to_pdf_via_word(outs)


if __name__ == "__main__":
    main()
