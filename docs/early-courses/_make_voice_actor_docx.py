# -*- coding: utf-8 -*-
"""Build combined voice-actor Word brief for both trial lessons."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD = ROOT / "08-voice-actor-trials-combined.md"
OUT = ROOT / "08-voice-actor-trials-combined.docx"
BLUE = RGBColor(0x1F, 0x4E, 0x86)
WARM = RGBColor(0x3A, 0x34, 0x2C)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def add_runs(paragraph, text, size=11):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 10, color=BLUE)
            continue
        segs = re.split(r"(\*\*[^*]+\*\*)", part)
        for seg in segs:
            if not seg:
                continue
            if seg.startswith("**") and seg.endswith("**"):
                run = paragraph.add_run(seg[2:-2])
                set_run_font(run, size=size, bold=True, color=WARM)
            else:
                run = paragraph.add_run(seg)
                set_run_font(run, size=size, color=WARM)


def add_table(doc, rows, widths=None, header_fill="1F4E86"):
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            txt = row[c_idx] if c_idx < len(row) else ""
            add_runs(p, txt, size=10.5 if r_idx else 10)
            if r_idx == 0:
                shade_cell(cell, header_fill)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            elif r_idx % 2 == 0:
                shade_cell(cell, "F7F1E8")
    if widths:
        set_cell_widths(table, widths)
    doc.add_paragraph()
    return table


def md_to_docx(md_path: Path, out_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = WARM

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            text = line[2:].strip()
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                set_run_font(run, size=20 if first_title else 16, bold=True, color=BLUE)
            if not first_title and text.startswith("Часть 2"):
                h.paragraph_format.page_break_before = True
            first_title = False
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                set_run_font(run, size=14, bold=True, color=BLUE)
        elif line.startswith("|") and "|" in line[1:]:
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in row):
                    i += 1
                    continue
                rows.append(row)
                i += 1
            widths = None
            if rows and len(rows[0]) == 4:
                widths = [1.2, 3.6, 3.4, 8.0]
            elif rows and len(rows[0]) == 3:
                widths = [1.2, 3.8, 11.2]
            elif rows and len(rows[0]) == 2:
                widths = [4.2, 12.0]
            add_table(doc, rows, widths)
            continue
        elif line.strip() == "---":
            pass
        elif line.strip().startswith("- ["):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line.strip()[2:])
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line.strip()[2:])
        elif re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", line.strip()))
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    doc.save(out_path)
    return out_path


def build():
    path = md_to_docx(MD, OUT)
    print(path)


if __name__ == "__main__":
    build()
